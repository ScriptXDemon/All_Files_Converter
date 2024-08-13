import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import img2pdf
import pypandoc
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx2txt import process
import nbformat
from nbconvert import HTMLExporter
import pdfkit
import tempfile
import os

st.title("Hari Files Converter")

# CSS for styling
st.markdown("""
    <style>
        .stButton>button {
            background: linear-gradient(135deg, #42a5f5, #4CAF50);
            color: white;
            padding: 10px 20px;
            margin: 8px 0;
            border: none;
            border-radius: 12px;
            cursor: pointer;
            width: 100%;
            font-size: 16px;
            font-weight: bold;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            transition: all 0.3s ease;
        }
        .stButton>button:hover {
            background: linear-gradient(135deg, #388E3C, #1E88E5);
            box-shadow: 0 6px 8px rgba(0, 0, 0, 0.15);
            transform: translateY(-2px);
        }
        .stButton>button:active {
            background: linear-gradient(135deg, #4CAF50, #42a5f5);
            box-shadow: 0 3px 4px rgba(0, 0, 0, 0.2);
            transform: translateY(0px);
        }
    </style>
    """, unsafe_allow_html=True)


# Conversion functions

def image_to_pdf(uploaded_file):
    # Load the image from the uploaded file
    image = Image.open(uploaded_file)
    
    # Seek to the beginning of the uploaded file to get the binary data
    uploaded_file.seek(0)
    
    # Convert the image to a PDF using the image binary data
    pdf_bytes = img2pdf.convert(uploaded_file.read())
    
    return pdf_bytes

def docx_to_html(docx_path):
    document = Document(docx_path)
    html_content = "<html><body>"

    for para in document.paragraphs:
        style = para.style
        font_size = None
        if style.font and style.font.size:
            font_size = style.font.size.pt
        else:
            font_size = Pt(12).pt  # Default font size

        html_content += f"<p style='font-size: {font_size}px;'>{para.text}</p>"

    html_content += "</body></html>"
    return html_content

def docx_to_pdf_func(uploaded_file):
    # Save the uploaded DOCX file to a temporary location
    temp_docx_path = "temp.docx"
    with open(temp_docx_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Convert DOCX to HTML
    html_content = docx_to_html(temp_docx_path)

    # Convert HTML to PDF using pdfkit
    output_pdf = "output.pdf"
    pdfkit.from_string(html_content, output_pdf)

    # Read the generated PDF file
    with open(output_pdf, "rb") as f:
        pdf_data = f.read()

    # Clean up temporary files
    os.remove(temp_docx_path)
    os.remove(output_pdf)

    return pdf_data

def ipynb_to_pdf(uploaded_file):
    notebook_content = uploaded_file.read().decode('utf-8')
    notebook = nbformat.reads(notebook_content, as_version=4)
    html_exporter = HTMLExporter()
    html_exporter.template_name = 'classic'  
    html_data, _ = html_exporter.from_notebook_node(notebook)
    pdf_data = pdfkit.from_string(html_data, False, options={"dpi": "300"})
    return pdf_data

def docx_to_image(uploaded_file):
    document = Document(uploaded_file)
    text = "\n".join([para.text for para in document.paragraphs])
    image = Image.new('RGB', (1600, 1200), color=(255, 255, 255))
    d = ImageDraw.Draw(image)
    d.text((10, 10), text, fill=(0, 0, 0))
    output_image = BytesIO()
    image.save(output_image, format="PNG", dpi=(300, 300))
    return output_image.getvalue()

def image_to_docx(uploaded_file):
    document = Document()
    image = Image.open(uploaded_file)
    max_width = Inches(6)
    width, height = image.size
    aspect_ratio = height / width
    document.add_picture(uploaded_file, width=max_width, height=max_width * aspect_ratio)
    output_docx = BytesIO()
    document.save(output_docx)
    return output_docx.getvalue()

# File uploader
uploaded_file = st.file_uploader("Choose a file to convert", type=["jpg", "jpeg", "png", "docx", "ipynb", "html"])

if uploaded_file is not None:
    # Select conversion type
    option = st.selectbox("Choose conversion type", 
                          ["Image to PDF", "DOCX to PDF", "IPYNB to PDF", "DOCX to Image", "Image to DOCX"])
    
    # Ask for confirmation before converting
    if st.checkbox("I confirm that I want to convert this file"):
        # Process the file based on the selected option
        if option == "Image to PDF" and uploaded_file.type.startswith('image/'):
            output_data = image_to_pdf(uploaded_file)
            output_mime = "application/pdf"
            output_ext = "pdf"
        elif option == "DOCX to PDF" and uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            output_data = docx_to_pdf_func(uploaded_file)
            output_mime = "application/pdf"
            output_ext = "pdf"
        elif option == "IPYNB to PDF" and uploaded_file.name.endswith('.ipynb'):
            output_data = ipynb_to_pdf(uploaded_file)
            output_mime = "application/pdf"
            output_ext = "pdf"
        elif option == "DOCX to Image" and uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            output_data = docx_to_image(uploaded_file)
            output_mime = "image/png"
            output_ext = "png"
        elif option == "Image to DOCX" and uploaded_file.type.startswith('image/'):
            output_data = image_to_docx(uploaded_file)
            output_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            output_ext = "docx"
        else:
            st.error("Unsupported file type for the selected conversion")
        
        if 'output_data' in locals():
            st.download_button(
                label="Download Converted File",
                data=output_data,
                file_name=f"output.{output_ext}",
                mime=output_mime
            )
    else:
        st.warning("Please confirm that you want to convert the file.")
